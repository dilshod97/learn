"""TXT va DOCX fayllardan matn ajratish."""
from pathlib import Path

ALLOWED_EXT = {".txt", ".docx"}


def read_file_text(path: str) -> str:
    p = Path(path)
    ext = p.suffix.lower()

    if ext == ".txt":
        # encoding aniq aniqlash
        for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
            try:
                return p.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        return p.read_bytes().decode("utf-8", errors="replace")

    if ext == ".docx":
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx o'rnatilmagan: pip install python-docx")
        doc = Document(str(p))
        parts = []
        # Paragraflar
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Jadvallar
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)

    raise ValueError(f"Qo'llab-quvvatlanmaydigan format: {ext} (faqat .txt va .docx)")
